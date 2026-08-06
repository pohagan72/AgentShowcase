"""Unit tests for features/pii_redaction/routes.py helpers.

Tests the redaction primitives directly (no HTTP, no GCS). We push an app
context because the helpers reference `flask.g.request_id` for logging;
without one, even the exception path can't finish (the log line uses `g`).
"""
from __future__ import annotations

import io
import warnings
from pathlib import Path

import pytest
from docx import Document

from features.pii_redaction.routes import (
    apply_redaction_to_text,
    build_analyzer,
    entities_for_mode,
    redact_runs_in_paragraph,
    redact_word_document_pii,
)


BLOCK = "█"
CV_FIXTURE = Path(__file__).parent / "test-files" / "alex_morgan_cv.docx"


@pytest.fixture(scope="module")
def analyzer():
    """The same ORG-enabled analyzer the running app uses. Slow to construct
    (~5s for the spaCy model) so we scope it to the module."""
    warnings.filterwarnings("ignore")
    try:
        return build_analyzer()
    except Exception as e:
        pytest.skip(f"build_analyzer() could not initialize: {e}")


@pytest.fixture(autouse=True)
def _push_app_context(app):
    with app.app_context():
        yield


class FakeResult:
    def __init__(self, start, end):
        self.start = start
        self.end = end


class TestApplyRedactionToText:
    def test_replaces_span_with_block_char(self):
        assert apply_redaction_to_text("hello world", [FakeResult(6, 11)]) == "hello " + BLOCK * 5

    def test_no_results_returns_original(self):
        assert apply_redaction_to_text("hello", []) == "hello"

    def test_multiple_non_overlapping(self):
        out = apply_redaction_to_text("ab cd ef", [FakeResult(0, 2), FakeResult(6, 8)])
        assert out == BLOCK * 2 + " cd " + BLOCK * 2

    def test_clamps_out_of_range_indices(self):
        assert apply_redaction_to_text("hi", [FakeResult(-5, 100)]) == BLOCK * 2


class TestRedactRunsInParagraph:
    """Exercise the multi-run overlap logic — the trickiest bit of the module."""

    def _para_with_runs(self, run_texts):
        doc = Document()
        p = doc.add_paragraph()
        for t in run_texts:
            p.add_run(t)
        return p

    def test_pii_within_single_run(self, analyzer):
        # Lowercase surrounding text avoids Presidio's NER tagging title-cased
        # words as PERSON, which would blank the anchor text and defeat the
        # "surrounding text survives" check.
        p = self._para_with_runs(["reach me at ", "john.doe@example.com", " thanks"])
        changed = redact_runs_in_paragraph(p, analyzer)
        assert changed
        assert "john.doe@example.com" not in p.text
        assert BLOCK in p.text
        assert "reach me at " in p.text
        assert " thanks" in p.text

    def test_pii_spanning_multiple_runs(self, analyzer):
        # Email split across two runs — the offset-mapping bug-prone case.
        p = self._para_with_runs(["reach john.doe", "@example.com now"])
        changed = redact_runs_in_paragraph(p, analyzer)
        assert changed
        assert "john.doe@example.com" not in p.text
        assert "reach " in p.text
        assert " now" in p.text

    def test_empty_paragraph_returns_false(self, analyzer):
        p = self._para_with_runs(["   "])
        assert redact_runs_in_paragraph(p, analyzer) is False

    def test_no_pii_returns_false(self, analyzer):
        p = self._para_with_runs(["The quick brown fox jumps."])
        assert redact_runs_in_paragraph(p, analyzer) is False
        assert p.text == "The quick brown fox jumps."

    def test_empty_run_between_pii_runs(self, analyzer):
        # python-docx will happily emit zero-length runs; the offset math must
        # skip them without going out of sync.
        p = self._para_with_runs(["Email ", "", "john.doe@example.com", " ok"])
        changed = redact_runs_in_paragraph(p, analyzer)
        assert changed
        assert "john.doe@example.com" not in p.text


class TestRedactWordDocument:
    """End-to-end at the document-object level — exercises paragraphs + tables."""

    def _build_doc_bytes(self, paragraphs, table_cells=None):
        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        if table_cells:
            table = doc.add_table(rows=len(table_cells), cols=len(table_cells[0]))
            for r, row in enumerate(table_cells):
                for c, cell_text in enumerate(row):
                    table.rows[r].cells[c].text = cell_text
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def _extract_paragraph_text(self, stream):
        stream.seek(0)
        doc = Document(stream)
        return [p.text for p in doc.paragraphs]

    def _extract_table_text(self, stream):
        stream.seek(0)
        doc = Document(stream)
        return [
            [cell.text for cell in row.cells]
            for table in doc.tables
            for row in table.rows
        ]

    def test_redacts_paragraph_pii(self, analyzer):
        src = self._build_doc_bytes([
            "Contact john.doe@example.com about the invoice.",
            "The quick brown fox jumps over the lazy dog.",
        ])
        out = redact_word_document_pii(src, analyzer)
        assert out is not None
        paras = self._extract_paragraph_text(out)
        assert "john.doe@example.com" not in paras[0]
        assert BLOCK in paras[0]
        assert paras[1] == "The quick brown fox jumps over the lazy dog."

    def test_redacts_table_cell_pii(self, analyzer):
        # Header cells use lowercase so Presidio's PERSON recognizer doesn't
        # tag them as names (it hits title-cased tokens aggressively).
        src = self._build_doc_bytes(
            paragraphs=["header row below:"],
            table_cells=[
                ["field", "value"],
                ["employee", "john.doe@example.com"],
            ],
        )
        out = redact_word_document_pii(src, analyzer)
        assert out is not None
        cells = self._extract_table_text(out)
        assert cells[0] == ["field", "value"]
        assert "john.doe@example.com" not in cells[1][1]
        assert BLOCK in cells[1][1]

    def test_returns_none_on_corrupt_input(self, analyzer):
        # Non-zip bytes will make python-docx raise; the handler is documented
        # to catch and return None so the route can surface a flash message.
        out = redact_word_document_pii(io.BytesIO(b"not a docx"), analyzer)
        assert out is None

    def test_seeded_sample_pii_gets_redacted(self, analyzer):
        """The full seed set from build_redact_sample.py — canary for regressions
        in Presidio version bumps or the run-offset math. Emails, phones, names,
        SSN, passport, and credit-card should all end up blanked."""
        src = self._build_doc_bytes([
            "Full name: John Doe",
            "Personal email: john.doe@example.com",
            "Mobile: (555) 123-4567",
            "The employee's tax identifier on file is 211-61-2524.",
            "Travel documentation: passport C12345678 (issued 2019).",
            "Card on file: Visa 4111-1111-1111-1111, expires 12/2028.",
            "Jane Doe (spouse) — (555) 987-6543, jane.doe@example.com",
        ])
        out = redact_word_document_pii(src, analyzer)
        assert out is not None
        joined = "\n".join(self._extract_paragraph_text(out))
        seeded = [
            "John Doe",
            "Jane Doe",
            "john.doe@example.com",
            "jane.doe@example.com",
            "(555) 123-4567",
            "(555) 987-6543",
            "211-61-2524",
            "C12345678",
            "4111-1111-1111-1111",
        ]
        survivors = [s for s in seeded if s in joined]
        assert not survivors, f"seeded PII survived redaction: {survivors}"


class TestEntitiesForMode:
    def test_unknown_mode_falls_back_to_pii_only(self):
        # Defensive: a missing / mangled form field must not accidentally trip
        # the more aggressive full-anon pass.
        assert entities_for_mode("garbage") == entities_for_mode("pii_only")
        assert entities_for_mode(None) == entities_for_mode("pii_only")

    def test_full_anon_is_superset_of_pii_only(self):
        pii = set(entities_for_mode("pii_only"))
        full = set(entities_for_mode("full_anon"))
        assert pii.issubset(full)
        # Full-anon must at minimum add LOCATION, DATE_TIME, and ORGANIZATION.
        # ORG only works because build_analyzer() re-enables it on the
        # SpacyRecognizer — Presidio's stock English config strips it out.
        assert {"LOCATION", "DATE_TIME", "ORGANIZATION"}.issubset(full - pii)


class TestModeContrast:
    """The whole point of the mode toggle: PII-only leaves orgs / dates alone,
    full-anon blanks them. If these diverge in a Presidio version bump we want
    to know immediately."""

    def _para(self, text):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run(text)
        return p

    def test_location_survives_pii_only_but_blanked_in_full_anon(self, analyzer):
        # "Manchester" was one of the CV cities the user didn't want blanked in
        # PII-only mode. Presidio's LOCATION recognizer catches it reliably.
        p1 = self._para("Operations Coordinator based in Manchester on the delivery team.")
        redact_runs_in_paragraph(p1, analyzer, entities=entities_for_mode("pii_only"))
        assert "Manchester" in p1.text

        p2 = self._para("Operations Coordinator based in Manchester on the delivery team.")
        redact_runs_in_paragraph(p2, analyzer, entities=entities_for_mode("full_anon"))
        assert "Manchester" not in p2.text
        assert BLOCK in p2.text

    def test_date_survives_pii_only_but_blanked_in_full_anon(self, analyzer):
        # "five years" is exactly the DATE_TIME hit we saw over-redact in the
        # Alex Morgan CV. In PII-only mode it must survive.
        p1 = self._para("Business operations professional with five years of experience.")
        redact_runs_in_paragraph(p1, analyzer, entities=entities_for_mode("pii_only"))
        assert "five years" in p1.text

        p2 = self._para("Business operations professional with five years of experience.")
        redact_runs_in_paragraph(p2, analyzer, entities=entities_for_mode("full_anon"))
        assert "five years" not in p2.text

    def test_email_blanked_in_both_modes(self, analyzer):
        # Sanity: the PII entities themselves must fire in both modes.
        for mode in ("pii_only", "full_anon"):
            p = self._para(f"contact john.doe@example.com now ({mode})")
            redact_runs_in_paragraph(p, analyzer, entities=entities_for_mode(mode))
            assert "john.doe@example.com" not in p.text, f"email survived {mode}"

    def test_organisation_survives_pii_only_but_blanked_in_full_anon(self, analyzer):
        # The whole point of wiring build_analyzer(): ORG must fire in full-anon.
        p1 = self._para("Employed by Northbridge Digital Services on contract.")
        redact_runs_in_paragraph(p1, analyzer, entities=entities_for_mode("pii_only"))
        assert "Northbridge Digital Services" in p1.text

        p2 = self._para("Employed by Northbridge Digital Services on contract.")
        redact_runs_in_paragraph(p2, analyzer, entities=entities_for_mode("full_anon"))
        assert "Northbridge Digital Services" not in p2.text
        assert BLOCK in p2.text


class TestAlexMorganCV:
    """End-to-end against the actual CV file that triggered the bug report.
    Pins the behaviour a reader would see for each mode. These will get noisy
    if Presidio or spaCy is upgraded — that's the point: we want to see any
    shift in what gets blanked before it hits production."""

    def _extract_all_text(self, stream):
        stream.seek(0)
        d = Document(stream)
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        parts.append(p.text)
        return "\n".join(parts)

    def _redact(self, analyzer, mode):
        assert CV_FIXTURE.exists(), f"missing fixture: {CV_FIXTURE}"
        stream = io.BytesIO(CV_FIXTURE.read_bytes())
        out = redact_word_document_pii(stream, analyzer, entities=entities_for_mode(mode))
        assert out is not None
        return self._extract_all_text(out)

    def test_pii_only_blanks_identifying_details_and_keeps_context(self, analyzer):
        text = self._redact(analyzer, "pii_only")

        # Identifying details must be gone.
        gone = [
            "Alex Morgan",
            "alex.morgan.test@example.com",
            "07123 456789",
            "linkedin.com/in/alex-morgan-test",
        ]
        for s in gone:
            assert s not in text, f"PII-only failed to blank {s!r}"

        # Context that PII-only mode must NOT touch — this is what the bug
        # report was about: Presidio's default was blanking these.
        for s in ["Manchester", "five years", "Microsoft", "April 2023", "Northbridge Digital Services"]:
            assert s in text, f"PII-only mode unexpectedly blanked {s!r}"

    def test_full_anon_blanks_orgs_locations_and_dates(self, analyzer):
        text = self._redact(analyzer, "full_anon")

        # Everything PII-only blanks is still gone.
        for s in ["Alex Morgan", "alex.morgan.test@example.com", "07123 456789"]:
            assert s not in text, f"full-anon failed on PII: {s!r}"

        # The extra stuff full-anon adds: cities, employers, dates.
        for s in [
            "Manchester",
            "Leeds",
            "Liverpool",
            "Northbridge Digital Services",
            "Greenfield Advisory Group",
            "Harbour Retail Solutions",
            "Fictional Northern University",
            "April 2023",
            "September 2018",
        ]:
            assert s not in text, f"full-anon failed to blank {s!r}"
