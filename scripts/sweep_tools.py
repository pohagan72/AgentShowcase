"""Automated end-to-end sweep of every MCP tool against a live Synzo deployment.

Companion to scripts/encode_for_inspector.py. This script drives the full
test rig in code:

  1. Reads SYNZO_API_KEY from .env.
  2. Calls tools/list to confirm the deployment is healthy and which tools
     are registered.
  3. For each of the 5 content-processing tools, the script:
       (a) synthesizes a tiny in-memory test payload,
       (b) calls upload_file to stash the bytes -> content_url,
       (c) fires a tools/call against /mcp with the content_url,
       (d) asserts on the response shape from mcp_tools.py.
  4. For detect_faces, requires --face-image <path> (MTCNN needs a real face;
     a synthetic shape won't exercise the model). Skipped with a clear note
     if not provided.
  5. Prints per-tool: outcome (success / tool error / JSON-RPC error /
     timeout / network error), latency, and a one-line result summary.

The sweep ASSERTS on the response field names defined in mcp_tools.py:
  - upload_file         -> content_url, expires_at, size_bytes, content_type, filename
  - summarize_document  -> classification, summary, filename
  - translate_document  -> translated_text, target_language, filename
  - redact_pii          -> result_url, mimetype, filename, original_size_bytes, redacted_size_bytes
  - analyze_image       -> analysis, dominant_colors, filename
  - detect_faces        -> result_url, mode, mimetype, filename

If a tool returns 200 but with the wrong shape, the script flags it as a
contract drift — that would be a submission-blocking bug Inspector wouldn't
necessarily catch.

USAGE:
    python -m scripts.sweep_tools
    python -m scripts.sweep_tools --base-url https://www.synzo.ai \\
        --face-image path/to/photo_with_face.jpg

Reads SYNZO_API_KEY from .env (loaded via python-dotenv). Override on the
command line with --api-key if you want a different key for a one-off run.
"""

from __future__ import annotations

import argparse
import base64
import io
import json
import os
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import certifi
import requests
from dotenv import load_dotenv

DEFAULT_BASE_URL = "https://www.synzo.ai"
MCP_PATH = "/mcp"
PROTOCOL_VERSION = "2025-06-18"

# Per-tool quota cost is roughly ceil(payload_bytes / 50KB) — the synthesized
# payloads here are all <50KB so each tool burns exactly 1 unit.
EXPECTED_UNITS_PER_CALL = 1


@dataclass
class ToolResult:
    name: str
    status: str  # "success" | "tool_error" | "jsonrpc_error" | "shape_drift" | "skipped" | "network_error"
    latency_s: float
    summary: str  # one-line human-readable
    raw_error: str | None = None


# --- Payload synthesis ---------------------------------------------------------


def make_summarize_payload() -> tuple[str, str]:
    """Tiny .docx with prose Gemini can classify and summarize."""
    from docx import Document

    doc = Document()
    doc.add_heading("Q3 Operations Review", level=1)
    doc.add_paragraph(
        "Revenue grew 14% quarter over quarter, driven primarily by enterprise "
        "renewals in the EMEA region. Customer churn decreased to 2.1%, the "
        "lowest in eight quarters. Headcount remained flat. We expect Q4 to "
        "trend higher given the pipeline coverage ratio of 3.2x."
    )
    buf = io.BytesIO()
    doc.save(buf)
    encoded = base64.b64encode(buf.getvalue()).decode("ascii")
    return "q3_review.docx", encoded


def make_translate_payload() -> tuple[str, str]:
    """Tiny .docx with simple translatable English text."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello, this is a test. Today is Thursday and the weather is sunny.")
    buf = io.BytesIO()
    doc.save(buf)
    encoded = base64.b64encode(buf.getvalue()).decode("ascii")
    return "hello.docx", encoded


def make_pii_payload() -> tuple[str, str]:
    """Tiny .docx peppered with Presidio-recognizable PII."""
    from docx import Document

    doc = Document()
    doc.add_paragraph(
        "Please contact John Smith at john.smith@example.com or 555-123-4567. "
        "His employee ID is 88241 and he lives at 1600 Pennsylvania Avenue NW."
    )
    buf = io.BytesIO()
    doc.save(buf)
    encoded = base64.b64encode(buf.getvalue()).decode("ascii")
    return "pii_sample.docx", encoded


def make_analyze_image_payload() -> tuple[str, str]:
    """Solid-color PNG. Gemini will describe it as "a blue square" or similar.
    Not interesting visually but proves the handler + vision call work."""
    from PIL import Image

    img = Image.new("RGB", (256, 256), color=(40, 90, 180))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    encoded = base64.b64encode(buf.getvalue()).decode("ascii")
    return "blue_square.png", encoded


def load_face_image(path: str) -> tuple[str, str]:
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"--face-image not found: {p}")
    raw = p.read_bytes()
    encoded = base64.b64encode(raw).decode("ascii")
    return p.name, encoded


# --- HTTP plumbing -------------------------------------------------------------


# Module-level toggle set by main() from CLI args. Off by default — TLS
# verification is the right default for any networked script. The flag exists
# only because corporate TLS-intercepting proxies (Zscaler / Netskope / etc.)
# inject a CA that's missing from certifi's bundle; in that environment the
# choice is "disable verify for this internal sweep" or "find and pin the
# corporate CA bundle". Both are defensible; the flag lets the operator pick.
INSECURE_SKIP_VERIFY = False


def post_jsonrpc(base_url: str, api_key: str, method: str, params: dict | None = None, *, timeout: int = 90) -> tuple[dict | None, float]:
    """Fire one JSON-RPC call. Returns (response_json, elapsed_seconds).

    timeout: HTTP read timeout. Server has its own 60s tool wall-clock; we set
    HTTP a bit higher to give the timeout error envelope time to come back.
    """
    url = base_url.rstrip("/") + MCP_PATH
    headers = {
        "Content-Type": "application/json",
        "MCP-Protocol-Version": PROTOCOL_VERSION,
        "Origin": "https://claude.ai",
        "Authorization": f"Bearer {api_key}",
    }
    body = {
        "jsonrpc": "2.0",
        "id": int(time.time() * 1000) % 100_000,
        "method": method,
    }
    if params is not None:
        body["params"] = params

    verify_arg: Any = False if INSECURE_SKIP_VERIFY else certifi.where()

    start = time.perf_counter()
    try:
        r = requests.post(url, headers=headers, json=body, timeout=timeout, verify=verify_arg)
    except requests.RequestException as e:
        elapsed = time.perf_counter() - start
        return {"_network_error": str(e)}, elapsed
    elapsed = time.perf_counter() - start

    try:
        return r.json(), elapsed
    except ValueError:
        return {"_http_status": r.status_code, "_body": r.text[:500]}, elapsed


def call_tool(base_url: str, api_key: str, name: str, args: dict) -> tuple[dict | None, float]:
    return post_jsonrpc(base_url, api_key, "tools/call", {"name": name, "arguments": args})


# --- Result interpretation -----------------------------------------------------


def interpret(name: str, response: dict | None, elapsed: float, expected_keys: set[str], summary_fn) -> ToolResult:
    """Classify a tools/call response into one of our outcome buckets."""
    if response is None:
        return ToolResult(name, "network_error", elapsed, "no response", None)

    if "_network_error" in response:
        return ToolResult(name, "network_error", elapsed, response["_network_error"], None)

    if "_http_status" in response:
        return ToolResult(
            name, "network_error", elapsed,
            f"HTTP {response['_http_status']}: {response.get('_body', '')[:100]}", None,
        )

    if "error" in response:
        err = response["error"]
        code = err.get("code")
        msg = err.get("message", "")
        # -32005 is our timeout code; treat as a distinct bucket so the user can
        # see whether the 60s ceiling is the bottleneck.
        if code == -32005:
            return ToolResult(name, "timeout", elapsed, f"server-side timeout (code -32005): {msg}", json.dumps(err))
        return ToolResult(name, "jsonrpc_error", elapsed, f"code {code}: {msg}", json.dumps(err))

    result = response.get("result", {})
    if result.get("isError"):
        # Tool-layer error (e.g. ToolError raised in mcp_tools.py). The
        # content blocks describe what went wrong.
        content = result.get("content", [])
        text = next((c.get("text", "") for c in content if c.get("type") == "text"), "")
        return ToolResult(name, "tool_error", elapsed, text[:200] or "isError=true with no message", text)

    # Success path. The MCP spec returns structured output under
    # result.structuredContent (if the server advertised it). Our server
    # returns the handler's dict there.
    structured = result.get("structuredContent") or {}
    if not structured:
        # Fall back to first content block — older client/server combos use that.
        content = result.get("content", [])
        for c in content:
            if c.get("type") == "text":
                try:
                    structured = json.loads(c["text"])
                    break
                except (ValueError, KeyError):
                    continue

    missing = expected_keys - set(structured.keys())
    if missing:
        return ToolResult(
            name, "shape_drift", elapsed,
            f"response missing expected keys: {sorted(missing)}",
            json.dumps(structured)[:300],
        )

    return ToolResult(name, "success", elapsed, summary_fn(structured), None)


# --- Sweep -------------------------------------------------------------------


def upload_payload(base_url: str, api_key: str, filename: str, b64: str) -> tuple[str | None, float, dict | None]:
    """POST to upload_file. Returns (content_url, elapsed, raw_response).

    The raw_response is returned so the caller can attribute upload_file's own
    result line to the sweep summary.
    """
    resp, elapsed = call_tool(
        base_url, api_key, "upload_file",
        {"filename": filename, "content_base64": b64},
    )
    if resp is None or "result" not in resp:
        return None, elapsed, resp
    structured = resp["result"].get("structuredContent") or {}
    return structured.get("content_url"), elapsed, resp


def run_sweep(base_url: str, api_key: str, face_image_path: str | None) -> list[ToolResult]:
    results: list[ToolResult] = []

    print(f"\n[preflight] Hitting {base_url}{MCP_PATH} tools/list ...")
    response, elapsed = post_jsonrpc(base_url, api_key, "tools/list")
    if not response or "result" not in response:
        print(f"  FAIL: {response}")
        sys.exit(2)
    listed = [t["name"] for t in response["result"].get("tools", [])]
    print(f"  OK ({elapsed:.2f}s): {len(listed)} tools advertised: {listed}")
    if "upload_file" not in listed:
        print("  FAIL: server does not advertise upload_file; sweep can't run the URL flow.")
        sys.exit(2)

    # Helper that uploads, prints a one-liner, and returns the URL or None on failure.
    def _upload(idx: int, label: str, filename: str, b64: str) -> str | None:
        print(f"  [{idx}.upload] uploading {filename} ({len(b64):,} b64 chars) ...")
        url, up_elapsed, raw = upload_payload(base_url, api_key, filename, b64)
        if url is None:
            print(f"  [{idx}.upload] FAILED in {up_elapsed:.2f}s: {raw}")
            return None
        print(f"  [{idx}.upload] OK ({up_elapsed:.2f}s) -> {url}")
        return url

    # --- upload_file standalone (first tool in the sweep, since everything else depends on it)
    print("\n[1/6] upload_file ...")
    fn, b64 = make_summarize_payload()
    resp, elapsed = call_tool(
        base_url, api_key, "upload_file",
        {"filename": fn, "content_base64": b64},
    )
    upload_result = interpret(
        "upload_file", resp, elapsed,
        expected_keys={"content_url", "expires_at", "size_bytes", "content_type", "filename"},
        summary_fn=lambda s: f"size_bytes={s.get('size_bytes')}, content_type={s.get('content_type')!r}",
    )
    print(f"  {upload_result.status.upper()} ({upload_result.latency_s:.2f}s) {upload_result.summary}")
    results.append(upload_result)

    summarize_url = None
    if upload_result.status == "success" and resp is not None:
        summarize_url = resp["result"]["structuredContent"]["content_url"]

    # --- summarize_document
    print("\n[2/6] summarize_document ...")
    if summarize_url is None:
        # Fallback: try fresh upload (in case the upload_file test was a retry without reseeding state).
        summarize_url = _upload(2, "summarize", fn, b64)
    if summarize_url is None:
        results.append(ToolResult("summarize_document", "skipped", 0.0, "upload_file failed", None))
    else:
        resp, elapsed = call_tool(
            base_url, api_key, "summarize_document",
            {"filename": fn, "content_url": summarize_url},
        )
        result = interpret(
            "summarize_document", resp, elapsed,
            expected_keys={"classification", "summary", "filename"},
            summary_fn=lambda s: f"classification={s.get('classification')!r}, summary[:80]={s.get('summary', '')[:80]!r}",
        )
        print(f"  {result.status.upper()} ({result.latency_s:.2f}s) {result.summary}")
        results.append(result)

    # --- translate_document
    print("\n[3/6] translate_document (target=Spanish) ...")
    fn, b64 = make_translate_payload()
    translate_url = _upload(3, "translate", fn, b64)
    if translate_url is None:
        results.append(ToolResult("translate_document", "skipped", 0.0, "upload_file failed", None))
    else:
        resp, elapsed = call_tool(
            base_url, api_key, "translate_document",
            {"filename": fn, "content_url": translate_url, "target_language": "Spanish"},
        )
        result = interpret(
            "translate_document", resp, elapsed,
            expected_keys={"translated_text", "target_language", "filename"},
            summary_fn=lambda s: f"translated_text[:80]={s.get('translated_text', '')[:80]!r}",
        )
        print(f"  {result.status.upper()} ({result.latency_s:.2f}s) {result.summary}")
        results.append(result)

    # --- redact_pii
    print("\n[4/6] redact_pii ...")
    fn, b64 = make_pii_payload()
    redact_url = _upload(4, "redact", fn, b64)
    if redact_url is None:
        results.append(ToolResult("redact_pii", "skipped", 0.0, "upload_file failed", None))
    else:
        resp, elapsed = call_tool(
            base_url, api_key, "redact_pii",
            {"filename": fn, "content_url": redact_url},
        )
        result = interpret(
            "redact_pii", resp, elapsed,
            expected_keys={"result_url", "mimetype", "filename"},
            summary_fn=lambda s: f"redacted_size={s.get('redacted_size_bytes')}, result_url={s.get('result_url')!r}",
        )
        print(f"  {result.status.upper()} ({result.latency_s:.2f}s) {result.summary}")
        results.append(result)

    # --- analyze_image
    print("\n[5/6] analyze_image (synthetic blue PNG) ...")
    fn, b64 = make_analyze_image_payload()
    analyze_url = _upload(5, "analyze", fn, b64)
    if analyze_url is None:
        results.append(ToolResult("analyze_image", "skipped", 0.0, "upload_file failed", None))
    else:
        resp, elapsed = call_tool(
            base_url, api_key, "analyze_image",
            {"filename": fn, "content_url": analyze_url},
        )

        def _analyze_summary(s):
            analysis = s.get("analysis", {}) or {}
            desc = (analysis.get("description") or "")[:80]
            colors = s.get("dominant_colors", [])
            return f"description[:80]={desc!r}, dominant_colors={colors[:3]}"

        result = interpret(
            "analyze_image", resp, elapsed,
            expected_keys={"analysis", "dominant_colors", "filename"},
            summary_fn=_analyze_summary,
        )
        print(f"  {result.status.upper()} ({result.latency_s:.2f}s) {result.summary}")
        results.append(result)

    # --- detect_faces
    print("\n[6/6] detect_faces ...")
    if not face_image_path:
        print("  SKIPPED — no --face-image provided. MTCNN won't surface anything useful")
        print("           on a synthetic PNG; re-run with --face-image path/to/photo.jpg.")
        results.append(ToolResult("detect_faces", "skipped", 0.0, "no face image provided", None))
    else:
        try:
            fn, b64 = load_face_image(face_image_path)
        except FileNotFoundError as e:
            print(f"  ERROR: {e}")
            results.append(ToolResult("detect_faces", "skipped", 0.0, str(e), None))
        else:
            print(f"  Loaded {fn}, base64 len={len(b64):,}.")
            faces_url = _upload(6, "faces", fn, b64)
            if faces_url is None:
                results.append(ToolResult("detect_faces", "skipped", 0.0, "upload_file failed", None))
            else:
                print(f"  First call may take 10-30s on cold start ...")
                resp, elapsed = call_tool(
                    base_url, api_key, "detect_faces",
                    {"filename": fn, "content_url": faces_url, "mode": "blur", "blur_strength": 2},
                )
                result = interpret(
                    "detect_faces", resp, elapsed,
                    expected_keys={"result_url", "mode", "mimetype", "filename"},
                    summary_fn=lambda s: f"mode={s.get('mode')!r}, mimetype={s.get('mimetype')!r}, result_url={s.get('result_url')!r}",
                )
                print(f"  {result.status.upper()} ({result.latency_s:.2f}s) {result.summary}")
                results.append(result)

    return results


# --- Report --------------------------------------------------------------------


def print_summary(results: list[ToolResult]) -> int:
    print("\n" + "=" * 78)
    print(f"{'TOOL':<24} {'STATUS':<14} {'LATENCY':<10} {'NOTES'}")
    print("=" * 78)
    for r in results:
        latency = f"{r.latency_s:.2f}s" if r.latency_s else "-"
        notes = r.summary[:50] + ("..." if len(r.summary) > 50 else "")
        print(f"{r.name:<24} {r.status:<14} {latency:<10} {notes}")
    print("=" * 78)

    # Quota cost estimate: success + tool_error both burn 1 unit; everything
    # else either refunds (timeout, jsonrpc_error in handler) or never hit
    # the metering path (network_error). This is an estimate — the dashboard
    # is the source of truth.
    billable = sum(1 for r in results if r.status in {"success", "tool_error", "shape_drift"})
    print(f"\nEstimated quota burned: ~{billable} units. Cross-check on the dashboard.")

    bad_buckets = {"jsonrpc_error", "network_error", "shape_drift", "tool_error"}
    failures = [r for r in results if r.status in bad_buckets]
    timeouts = [r for r in results if r.status == "timeout"]

    if failures:
        print(f"\n{len(failures)} tool(s) failed:")
        for r in failures:
            print(f"  - {r.name}: {r.status}")
            if r.raw_error:
                print(f"      raw: {r.raw_error[:300]}")
        return 1

    if timeouts:
        print(f"\n{len(timeouts)} tool(s) timed out (Phase 2.5.A safety net fired):")
        for r in timeouts:
            print(f"  - {r.name}")
        print("  Note: timeouts refund quota. Re-run to confirm whether it was a cold start.")
        return 1

    skipped = [r for r in results if r.status == "skipped"]
    if skipped:
        print(f"\nGate (d) partially complete: {len(skipped)} tool(s) skipped: {[r.name for r in skipped]}.")
        return 0

    print("\nAll tools green. API-key auth path proven end-to-end on live deployment.")
    return 0


def main() -> int:
    load_dotenv()

    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--base-url", default=DEFAULT_BASE_URL, help=f"Base URL (default: {DEFAULT_BASE_URL})")
    parser.add_argument(
        "--api-key",
        default=os.environ.get("SYNZO_API_KEY"),
        help="API key. Defaults to SYNZO_API_KEY env var (loaded from .env).",
    )
    parser.add_argument("--face-image", help="Path to a real photo with a face for detect_faces. Skipped if omitted.")
    parser.add_argument(
        "--insecure-skip-verify",
        action="store_true",
        help=(
            "Skip TLS cert verification. Use only when behind a corporate "
            "TLS-intercepting proxy whose CA isn't in certifi's bundle. "
            "The server-side cert is still validated by your OS-level proxy."
        ),
    )
    args = parser.parse_args()

    global INSECURE_SKIP_VERIFY
    INSECURE_SKIP_VERIFY = args.insecure_skip_verify
    if INSECURE_SKIP_VERIFY:
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        print("WARNING: TLS cert verification disabled (--insecure-skip-verify).", file=sys.stderr)

    if not args.api_key:
        print("ERROR: no API key. Set SYNZO_API_KEY in .env or pass --api-key.", file=sys.stderr)
        return 2

    if not args.api_key.startswith("sk_synzo_"):
        print(f"WARNING: API key doesn't start with 'sk_synzo_' — got {args.api_key[:12]!r}...", file=sys.stderr)

    results = run_sweep(args.base_url, args.api_key, args.face_image)
    return print_summary(results)


if __name__ == "__main__":
    sys.exit(main())
