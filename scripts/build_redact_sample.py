"""Generate redact-sample.docx for the Anthropic MCP submission reviewer bundle.

The document is a fake internal HR memo with synthetic PII seeded into the body.
Every value is fabricated and non-functional (the Visa is the standard 4111-...
test card; the SSN is a non-issuable range; the passport is a fake nine-char
string). No real person, account, or document is referenced.

The point is to give a reviewer something they can run through `redact_pii` and
visibly see Presidio find and block out the seeded entities. Detection only
needs Presidio's default English recognizers (PERSON, EMAIL_ADDRESS,
PHONE_NUMBER, US_SSN, CREDIT_CARD, US_PASSPORT) which the production handler
runs with no entities= filter.

Run from the repo root:
    .venv/Scripts/python scripts/build_redact_sample.py <output_path>
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


def build(output_path: Path) -> None:
    doc = Document()

    title = doc.add_heading("CONFIDENTIAL — Internal HR Memo", level=1)
    for run in title.runs:
        run.font.size = Pt(16)

    doc.add_paragraph(
        "From: HR Operations\n"
        "To: Payroll Team\n"
        "Re: New hire onboarding — record verification"
    )

    doc.add_paragraph(
        "Please process the onboarding records for our new hire, "
        "John Doe, who joins the engineering team on Monday. The records below "
        "have been provided for payroll setup and benefits enrollment."
    )

    doc.add_heading("Employee Information", level=2)
    # Address kept at city/state level — Presidio's default LOCATION recognizer
    # catches Springfield/IL but has no street-level address recognizer, so
    # including "1234 Maple Street" would leave a visible survivor in the demo.
    doc.add_paragraph(
        "Full name: John Doe\n"
        "Personal email: john.doe@example.com\n"
        "Mobile: (555) 123-4567\n"
        "Date of birth: March 14, 1990\n"
        "Location: Springfield, IL"
    )

    doc.add_heading("Identification & Tax Records", level=2)
    # Presidio's UsSsnRecognizer needs the SSN in cleaner context — putting it
    # on its own paragraph after a descriptive sentence helps the recognizer
    # latch onto the value rather than just the descriptive label. The same
    # applies to the passport on the next paragraph.
    doc.add_paragraph(
        "The employee's tax identifier on file is 211-61-2524. "
        "Please verify against the W-4 submitted in the onboarding portal."
    )
    doc.add_paragraph(
        "Travel documentation: passport C12345678 (issued 2019, expires 2029)."
    )

    doc.add_heading("Payroll Direct Deposit", level=2)
    doc.add_paragraph(
        "Card on file (for expense reimbursement only): "
        "Visa 4111-1111-1111-1111, expires 12/2028.\n"
        "Routing/account details will be collected separately via the "
        "secure payroll portal."
    )

    doc.add_heading("Emergency Contact", level=2)
    doc.add_paragraph(
        "Jane Doe (spouse) — (555) 987-6543, jane.doe@example.com"
    )

    doc.add_paragraph(
        "Please confirm receipt and flag any missing fields by Friday. "
        "Reach out to HR Operations directly if anything looks off."
    )

    doc.add_paragraph(
        "— HR Operations, Acme Corp"
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Wrote {output_path} ({output_path.stat().st_size} bytes)")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("usage: build_redact_sample.py <output_path>", file=sys.stderr)
        sys.exit(2)
    build(Path(sys.argv[1]))
